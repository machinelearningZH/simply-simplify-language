# ---------------------------------------------------------------
# Imports

import streamlit as st

st.set_page_config(layout="wide")

import io
import logging
import os
import time
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt
from dotenv import load_dotenv

from openai import OpenAI
from zix.understandability import get_cefr, get_zix

from app_core import (
    APP_DIR,
    ResultState,
    app_path,
    build_log_payload,
    classify_understandability,
    configure_event_logger,
    create_prompt,
    extract_tagged_response,
    format_one_click_results,
    format_understandability_message,
    load_project_info,
    load_yaml_config,
    repo_path,
    result_models_used,
    rounded_score,
    strip_markdown,
    write_event_log,
)
from utils_prompts import SAMPLE_TEXT

# ---------------------------------------------------------------
# Constants

load_dotenv(app_path(".env"))

logger = logging.getLogger(__name__)

API_KEYS = {
    "OPENROUTER": os.getenv("OPENROUTER_API_KEY"),
}


@st.cache_resource
def load_config():
    """Load configuration from YAML file."""
    return load_yaml_config(repo_path("config.yaml"))


# Load configuration
config = load_config()

# Create model dictionaries from config
MODEL_IDS = {model["name"]: model["id"] for model in config["models"]}
MODEL_NAMES = list(MODEL_IDS.keys())

# Get configuration values from config
TEMPERATURE = config["api"]["temperature"]
MAX_TOKENS = config["api"]["max_tokens"]
API_BASE_URL = config["api"]["base_url"]
API_TIMEOUT_SECONDS = config["api"]["timeout_seconds"]
API_MAX_RETRIES = config["api"]["max_retries"]
TEXT_AREA_HEIGHT = config["ui"]["text_area_height"]
MAX_CHARS_INPUT = config["ui"]["max_chars_input"]
USER_WARNING = f"<sub>{config['ui']['user_warning']}</sub>"

# Document formatting constants
FONT_WORDDOC = config["document"]["font_name"]
FONT_SIZE_HEADING = config["document"]["font_size_heading"]
FONT_SIZE_PARAGRAPH = config["document"]["font_size_paragraph"]
FONT_SIZE_FOOTER = config["document"]["font_size_footer"]
DEFAULT_OUTPUT_FILENAME = config["document"]["default_output_filename"]
ANALYSIS_FILENAME = config["document"]["analysis_filename"]
PAGE_WIDTH_INCHES = config["document"]["page_width_inches"]
PAGE_HEIGHT_INCHES = config["document"]["page_height_inches"]
DOWNLOAD_MIME_TYPE = config["document"]["download_mime_type"]

# Understandability limits
LIMIT_HARD = config["understandability"]["limit_hard"]
LIMIT_MEDIUM = config["understandability"]["limit_medium"]
METRIC_LABEL = config["understandability"]["metric_label"]
METRIC_HELP = config["understandability"]["metric_help"]

DATETIME_FORMAT = config["app"]["datetime_format"]
EVENT_LOGGER = configure_event_logger(config["logging"], base_dir=APP_DIR)

# ---------------------------------------------------------------
# Functions


@st.cache_resource
def get_project_info():
    """Get markdown for project information that is shown in the expander section at the top of the app."""
    return load_project_info()


def create_project_info(project_info):
    """Create expander for project info. Add the image in the middle of the content."""
    with st.expander("Detaillierte Informationen zum Projekt"):
        project_info = project_info.split("ADD_IMAGE_HERE")
        st.markdown(project_info[0], unsafe_allow_html=True)
        st.image(str(app_path("zix_scores.jpg")), width="content")
        st.markdown(project_info[1], unsafe_allow_html=True)


def get_result_from_response(response):
    """Extract text between tags from response."""
    tag = "leichtesprache" if leichte_sprache else "einfachesprache"
    return extract_tagged_response(response, tag)


@st.cache_resource
def get_openrouter_client():
    if not API_KEYS["OPENROUTER"]:
        raise ValueError("OPENROUTER_API_KEY is not set")
    return OpenAI(
        base_url=API_BASE_URL,
        api_key=API_KEYS["OPENROUTER"],
        timeout=API_TIMEOUT_SECONDS,
        max_retries=API_MAX_RETRIES,
    )


def invoke_model(
    text,
    model_id,
    analysis=False,
):
    """Invoke any model through OpenRouter."""
    final_prompt, system = create_prompt(
        text,
        analysis=analysis,
        leichte_sprache=leichte_sprache,
        condense_text=condense_text,
    )

    try:
        message = openrouter_client.chat.completions.create(
            model=model_id,
            temperature=TEMPERATURE,
            max_tokens=MAX_TOKENS,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": final_prompt},
            ],
        )
        content = message.choices[0].message.content
        if content is None:
            raise ValueError("No content received from API")

        message = content.strip()
        message = get_result_from_response(message)
        message = strip_markdown(message)
        return True, message
    except Exception:
        logger.exception("Model invocation failed for model_id=%s", model_id)
        return False, "Model response could not be created."


def enter_sample_text():
    """Enter sample text into the text input in the left column."""
    st.session_state.key_textinput = SAMPLE_TEXT


def get_one_click_results():
    with ThreadPoolExecutor(max_workers=len(MODEL_IDS)) as executor:
        futures = {
            name: executor.submit(
                invoke_model,
                st.session_state.key_textinput,
                model_id,
            )
            for name, model_id in MODEL_IDS.items()
        }

    # invoke_model never raises: it returns (False, message) on failure.
    responses = {name: future.result() for name, future in futures.items()}

    return format_one_click_results(responses, score_fn=get_zix, cefr_fn=get_cefr)


def create_download_link(result):
    """Create a downloadable Word document and download link of the results."""
    document = Document()

    h1 = document.add_heading("Ausgangstext")
    p1 = document.add_paragraph("\n" + result.source_text)

    if result.analysis:
        h2 = document.add_heading(f"Analyse von Sprachmodell {result.model_choice}")
    elif result.one_click:
        h2 = document.add_heading("Vereinfachte Texte von Sprachmodellen")
    else:
        h2 = document.add_heading("Vereinfachter Text von Sprachmodell")

    p2 = document.add_paragraph(result.response)

    timestamp = datetime.now().strftime(DATETIME_FORMAT)
    models_used = result_models_used(result)
    footer = document.sections[0].footer
    footer.paragraphs[
        0
    ].text = f"Erstellt am {timestamp} mit der Prototyp-App «Einfache Sprache», Statistisches Amt, Kanton Zürich.\nSprachmodell(e): {models_used}\nVerarbeitungszeit: {result.time_processed:.1f} Sekunden"

    # Set font for all paragraphs.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_WORDDOC

    # Set font size for all headings.
    for paragraph in [h1, h2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_HEADING)

    # Set font size for all paragraphs.
    for paragraph in [p1, p2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_PARAGRAPH)

    # Set font and font size for footer.
    for run in footer.paragraphs[0].runs:
        run.font.name = FONT_WORDDOC
        run.font.size = Pt(FONT_SIZE_FOOTER)

    section = document.sections[0]
    section.page_width = Inches(PAGE_WIDTH_INCHES)
    section.page_height = Inches(PAGE_HEIGHT_INCHES)

    io_stream = io.BytesIO()
    document.save(io_stream)

    file_name = DEFAULT_OUTPUT_FILENAME

    if result.one_click:
        caption = "Vereinfachte Texte herunterladen"
    else:
        caption = "Vereinfachten Text herunterladen"

    if result.analysis:
        file_name = ANALYSIS_FILENAME
        caption = "Analyse herunterladen"
    st.download_button(
        label=caption,
        data=io_stream.getvalue(),
        file_name=file_name,
        mime=DOWNLOAD_MIME_TYPE,
    )


def log_event(
    text,
    response,
    do_analysis,
    do_simplification,
    do_one_click,
    leichte_sprache,
    model_choice,
    time_processed,
    success,
):
    """Log event."""
    payload = build_log_payload(
        text=text,
        response=response,
        do_analysis=do_analysis,
        do_simplification=do_simplification,
        do_one_click=do_one_click,
        leichte_sprache=leichte_sprache,
        model_choice=model_choice,
        time_processed=time_processed,
        success=success,
        datetime_format=DATETIME_FORMAT,
    )
    write_event_log(EVENT_LOGGER, payload)


def render_download_and_caption(result):
    """Render the download button and processing-time caption for a result."""
    create_download_link(result)
    st.caption(f"Verarbeitet in {result.time_processed:.1f} Sekunden.")


def render_result(result):
    """Render the latest generated result from session state."""
    text = "Dein vereinfachter Text"
    if result.analysis:
        text = "Deine Analyse"

    with placeholder_result.container():
        st.text_area(
            text,
            height=TEXT_AREA_HEIGHT,
            value=result.response,
        )

        # One-click aggregates several models into one text. A single
        # understandability score for the concatenated output is not meaningful
        # (per-model scores are shown inline instead), so we show the source
        # score and offer the download.
        if result.one_click:
            with placeholder_analysis.container():
                st.metric(
                    label=METRIC_LABEL,
                    value=rounded_score(result.score_source),
                    help=METRIC_HELP,
                )
                render_download_and_caption(result)
        elif result.simplification:
            score_target = get_zix(result.response)
            score_target_rounded = rounded_score(score_target)
            cefr_target = get_cefr(score_target)
            target_classification = classify_understandability(
                score_target,
                limit_hard=LIMIT_HARD,
                limit_medium=LIMIT_MEDIUM,
            )
            st.markdown(
                format_understandability_message(
                    "vereinfachter Text",
                    score_target_rounded,
                    cefr_target,
                    target_classification,
                )
            )
            with placeholder_analysis.container():
                st.metric(
                    label=METRIC_LABEL,
                    value=score_target_rounded,
                    delta=rounded_score(score_target - result.score_source),
                    help=METRIC_HELP,
                )
                render_download_and_caption(result)
        else:
            with placeholder_analysis.container():
                st.metric(
                    label=METRIC_LABEL,
                    value=rounded_score(result.score_source),
                    help=METRIC_HELP,
                )
                render_download_and_caption(result)


# ---------------------------------------------------------------
# Main

try:
    openrouter_client = get_openrouter_client()
except ValueError:
    st.error(
        "OPENROUTER_API_KEY fehlt. Bitte hinterlege den API-Key in _streamlit_app/.env."
    )
    st.stop()

project_info = get_project_info()

# Persist text input across sessions in session state.
# Otherwise, the text input sometimes gets lost when the user clicks on a button.
if "key_textinput" not in st.session_state:
    st.session_state.key_textinput = ""

st.markdown("## 🙋‍♀️ Sprache einfach vereinfachen")
create_project_info(project_info)
st.caption(USER_WARNING, unsafe_allow_html=True)
st.markdown("---")

# Set up first row with all buttons and settings.
button_cols = st.columns([1, 1, 1, 2])
with button_cols[0]:
    st.button(
        "Beispiel einfügen",
        on_click=enter_sample_text,
        width="stretch",
        type="secondary",
        help="Fügt einen Beispieltext ein.",
    )
    do_analysis = st.button(
        "Analysieren",
        width="stretch",
        help="Analysiert deinen Ausgangstext Satz für Satz.",
    )
with button_cols[1]:
    do_simplification = st.button(
        "Vereinfachen",
        width="stretch",
        help="Vereinfacht deinen Ausgangstext.",
    )
    do_one_click = st.button(
        "🚀 One-Klick",
        width="stretch",
        help="Schickt deinen Ausgangstext gleichzeitig an alle Modelle.",
    )
with button_cols[2]:
    leichte_sprache = st.toggle(
        "Leichte Sprache",
        value=False,
        help="**Schalter aktiviert**: «Leichte Sprache». **Schalter nicht aktiviert**: «Einfache Sprache».",
    )
    condense_text = False
    if leichte_sprache:
        condense_text = st.toggle(
            "Text verdichten",
            value=True,
            help="**Schalter aktiviert**: Modell konzentriert sich auf essentielle Informationen und versucht, Unwichtiges wegzulassen. **Schalter nicht aktiviert**: Modell versucht, alle Informationen zu übernehmen.",
        )
with button_cols[3]:
    model_choice = st.radio(
        label="Sprachmodell",
        options=MODEL_NAMES,
        index=0,
        horizontal=True,
    )

# Instantiate empty containers for the text areas.
cols = st.columns([2, 2, 1])

with cols[0]:
    source_text = st.container()
with cols[1]:
    placeholder_result = st.empty()
with cols[2]:
    placeholder_analysis = st.empty()

# Populate containers.
with source_text:
    st.text_area(
        "Ausgangstext, den du vereinfachen möchtest",
        value=None,
        height=TEXT_AREA_HEIGHT,
        max_chars=MAX_CHARS_INPUT,
        key="key_textinput",
    )
with placeholder_result:
    st.text_area(
        "Ergebnis",
        height=TEXT_AREA_HEIGHT,
    )
with placeholder_analysis:
    st.metric(
        label=METRIC_LABEL,
        value=None,
        delta=None,
        help=METRIC_HELP,
    )


# Derive model_id from explicit model_choice.
model_id = MODEL_IDS[model_choice]

# Start processing if one of the processing buttons is clicked.
if do_simplification or do_analysis or do_one_click:
    start_time = time.time()
    if st.session_state.key_textinput == "":
        st.error("Bitte gib einen Text ein.")
        st.stop()

    score_source = get_zix(st.session_state.key_textinput)
    score_source_rounded = rounded_score(score_source)
    cefr_source = get_cefr(score_source)
    source_classification = classify_understandability(
        score_source,
        limit_hard=LIMIT_HARD,
        limit_medium=LIMIT_MEDIUM,
    )

    # Analyze source text and display results.
    with source_text:
        st.markdown(
            format_understandability_message(
                "Ausgangstext",
                score_source_rounded,
                cefr_source,
                source_classification,
            )
        )
        with placeholder_analysis.container():
            st.metric(
                label=METRIC_LABEL,
                value=score_source_rounded,
                delta=None,
                help=METRIC_HELP,
            )

        with placeholder_analysis.container():
            with st.spinner("Ich arbeite..."):
                # One-click simplification.
                if do_one_click:
                    success, response = get_one_click_results()
                # Regular text simplification or analysis
                else:
                    success, response = invoke_model(
                        st.session_state.key_textinput,
                        model_id=model_id,
                        analysis=do_analysis,
                    )

    if success is False:
        st.error(
            "Es ist ein Fehler bei der Abfrage der APIs aufgetreten. Bitte versuche es erneut. Alternativ überprüfe Code, API-Keys, Verfügbarkeit der Modelle und ggf. Internetverbindung."
        )
        time_processed = time.time() - start_time
        log_event(
            st.session_state.key_textinput,
            "Error from model call",
            do_analysis,
            do_simplification,
            do_one_click,
            leichte_sprache,
            model_choice,
            time_processed,
            success,
        )

        st.stop()

    # Often the models return the German letter «ß». Replace it with the Swiss «ss».
    response = response.replace("ß", "ss")
    time_processed = time.time() - start_time

    result = ResultState(
        source_text=st.session_state.key_textinput,
        response=response,
        analysis=do_analysis,
        simplification=do_simplification,
        one_click=do_one_click,
        model_choice=model_choice,
        model_names=tuple(MODEL_NAMES),
        time_processed=time_processed,
        score_source=score_source,
    )
    st.session_state.last_result = result
    render_result(result)

    log_event(
        st.session_state.key_textinput,
        response,
        do_analysis,
        do_simplification,
        do_one_click,
        leichte_sprache,
        model_choice,
        time_processed,
        success,
    )
    st.stop()

if "last_result" in st.session_state:
    render_result(st.session_state.last_result)
