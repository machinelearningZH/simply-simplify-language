# ---------------------------------------------------------------
# Imports

import streamlit as st

st.set_page_config(layout="wide")

import os
import re
from datetime import datetime
import time
import base64
from docx import Document
from docx.shared import Pt, Inches
import io
from dotenv import load_dotenv

import logging

logging.basicConfig(
    filename="app.log",
    datefmt="%d-%b-%y %H:%M:%S",
    level=logging.WARNING,
)

import pandas as pd
import numpy as np
import spacy
from spacy.language import Language
import textdescriptives as td

# For usage of the Azure OpenAI client, see: https://github.com/openai/openai-python?tab=readme-ov-file#microsoft-azure-openai
from openai import AzureOpenAI

from utils_sample_texts import (
    SAMPLE_TEXT_01,
)

from utils_prompts import (
    SYSTEM_MESSAGE_ES,
    SYSTEM_MESSAGE_LS,
    RULES_ES,
    RULES_LS,
    REWRITE_COMPLETE,
    REWRITE_CONDENSED,
    OPENAI_TEMPLATE_ES,
    OPENAI_TEMPLATE_LS,
    OPENAI_TEMPLATE_ANALYSIS_ES,
    OPENAI_TEMPLATE_ANALYSIS_LS,
)


OPENAI_TEMPLATES = [
    OPENAI_TEMPLATE_ES,
    OPENAI_TEMPLATE_LS,
    OPENAI_TEMPLATE_ANALYSIS_ES,
    OPENAI_TEMPLATE_ANALYSIS_LS,
]


# ---------------------------------------------------------------
# Constants

load_dotenv()

AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")

# The actually used model is defined by the 'deploymentName', see docs.
MODEL_CHOICE = f"Azure OpenAI ({AZURE_OPENAI_DEPLOYMENT})"
MODEL_NAME = "GPT-4o"

# From our testing we derive a sensible temperature of 0.5 as a good trade-off between creativity and coherence. Adjust this to your needs.
TEMPERATURE = 0.5

# Height of the text areas for input and output.
TEXT_AREA_HEIGHT = 600

# Maximum number of characters for the input text.
# This is way below the context window sizes of the models.
# Adjust to your needs. However, we found that users can work and validate better when we nudge to work with shorter texts.
MAX_CHARS_INPUT = 10_000


USER_WARNING = """<sub>⚠️ Achtung: Diese App ist ein Prototyp. Nutze die App :red[**nur für öffentliche, nicht sensible Daten**]. Die App liefert lediglich einen Textentwurf. Überprüfe das Ergebnis immer und passe es an, wenn nötig. Die aktuelle App-Version ist v0.1.1 Die letzte Aktualisierung war am 14.6.2024."""


# Constants for the formatting of the Word document that can be downloaded.
FONT_WORDDOC = "Arial"
FONT_SIZE_HEADING = 12
FONT_SIZE_PARAGRAPH = 9
FONT_SIZE_FOOTER = 7


# Limits for the understandability score to determine if the text is easy, medium or hard to understand.
LIMIT_HARD = 13
LIMIT_MEDIUM = 16


# ---------------------------------------------------------------
# Functions


@st.cache_resource
def get_project_info():
    """Get markdown for project information that is shown in the expander section at the top of the app."""
    with open("utils_expander_openai.md") as f:
        return f.read()


def create_project_info(project_info):
    """Create expander for project info. Add the image in the middle of the content."""
    with st.expander("Detaillierte Informationen zum Projekt"):
        project_info = project_info.split("### Image ###")
        st.markdown(project_info[0], unsafe_allow_html=True)
        st.image("score.png", use_column_width=True)
        st.markdown(project_info[1], unsafe_allow_html=True)


@st.cache_resource
def get_word_scores():
    """Get commond word scores from the parquet file.

    This is a list of common German words in texts written in Einfache and Leichte Sprache. We use this ranking to calculate, how common the vocabulary in the text ist and therefore how easy the text is to understand.

    We have lemmatized and lower cased the words. Also note that the German `ß` has been replaced with `ss`.
    """
    word_scores = pd.read_parquet("word_scores.parq")
    word_scores = dict(zip(word_scores["lemma"], word_scores["score"]))
    return word_scores


@st.cache_resource
def get_nlp_pipeline():
    """Get NLP pipeline.

    We create a spacy pipeline with a custom component to calculate the common word score. We also add the textdescriptives components for readability and descriptive statistics.
    """

    word_scores = get_word_scores()

    @Language.component("common_word_score")
    def extract_common_word_score(doc):
        """Extract common word score and add to doc.user_data."""

        # Calculate common word score.
        doc_len = len([x for x in doc if not x.is_punct and not x.like_num])
        doc_scores = 0
        for token in doc:
            lemma = token.lemma_.lower()
            if lemma in word_scores:
                doc_scores += word_scores[lemma]

        doc_scores = doc_scores / doc_len
        doc.user_data["common_word_score"] = doc_scores / 1000

        return doc

    nlp = spacy.load("de_core_news_sm")
    nlp.add_pipe("textdescriptives/descriptive_stats")
    nlp.add_pipe("textdescriptives/readability")
    nlp.add_pipe("common_word_score")
    return nlp


def extract_text_features(text):
    """Extract text features from text.

    We use the spacy pipeline to extract text features needed for the understandability formula and return them as a DataFrame.
    """
    doc = nlp(text)
    row_metrics = td.extractors.extract_dict(
        doc,
        metrics=[
            "descriptive_stats",
            "readability",
        ],
        include_text=False,
    )
    for metric in [
        "common_word_score",
    ]:
        row_metrics[0][metric] = doc.user_data[metric]
    return pd.DataFrame.from_dict(row_metrics)


def punctuate_paragraphs_and_bulleted_lists(text):
    """Add a dot to lines that do not end with a dot and do some additional clean up to correctly calculate the understandability score.

    Texts in Einfache and Leichte Sprache often do not end sentences with a dot. This function adds a dot to lines that do not end with a dot. It also removes bullet points and hyphens from the beginning of lines and removes all line breaks and multiple spaces. This is necessary to correctly calculate the understandability score.
    """

    # Add a space to the end of the text to properly process the last sentence.
    text = text + " "

    # This regex pattern matches lines that do not end with a dot and are not empty.
    pattern = r"(?<!(\.))[^\n]$"

    # Replace lines not ending with a dot with the same line plus a dot.
    # The 're.MULTILINE' flag treats each line as a separate string.
    new_text = re.sub(pattern, r"\g<0>.", text, flags=re.MULTILINE)

    # Remove bullet points and hyphens from the beginning of lines.
    new_text = re.sub(r"^[\s]*[-•]", "", new_text, flags=re.MULTILINE)

    # Remove all line breaks and multiple spaces.
    new_text = new_text.replace("\n", " ")
    new_text = re.sub(r"\s+", " ", new_text)

    return new_text.strip()


def calculate_understandability(data):
    """Calculate understandability score from text metrics.

    We derived this formula from a dataset of legal and administrative texts, as well as Einfache and Leichte Sprache. We trained a Logistic Regression model to differentiate between complex and simple texts. From the most significant model coefficients we devised this formula to estimate a text's understandability.

    It is not a perfect measure, but it gives a good indication of how easy a text is to understand.

    We do not take into account a couple of text features that are relevant to understandability, such as the use of passive voice, the use of pronouns, or the use of complex sentence structures. These features are not easily extracted from text and would require a more complex model to calculate.
    """
    cws = (data.common_word_score - 7.8) / 1.1
    rix = (data.rix - 3.9) / 1.7
    sls = (data.sentence_length_std - 6.4) / 4.2
    slm = (data.sentence_length_mean - 11.7) / 3.7
    cws = 1 - cws

    score = ((cws * 0.2 + rix * 0.325 + sls * 0.225 + slm * 0.15) + 1.3) * 3.5

    # We clip the score to a range of 0 to 20.
    score = 20 - (score.values[0])
    if score < 0:
        score = 0
    if score > 20:
        score = 20
    return score


def get_cefr_level(score):
    """Get CEFR level from understandability score.

    We calculated these ranges by scoring various text samples where we had an approximate idea of their CEFR level. Again these ranges are not perfect, but give a good indication of the CEFR level.
    """
    if score >= 18.3:
        return "A1"
    elif score >= 17.7 and score < 18.3:
        return "A2"
    elif score >= 17.5 and score < 17.7:
        return "A2 bis B1"
    elif score >= 15.7 and score < 17.5:
        return "B1"
    elif score >= 13.7 and score < 15.7:
        return "B2"
    elif score >= 12.4 and score < 13.7:
        return "C1"
    elif score >= 12.2 and score < 12.4:
        return "C1 bis C2"
    elif score < 12.2:
        return "C2"


def get_understandability(text):
    """Get understandability score from text."""
    # Replace newlines with punctuation, so that paragraphs
    # and bullet point lists are not counted as complex sentences.
    text = punctuate_paragraphs_and_bulleted_lists(text)

    features = extract_text_features(text)
    return calculate_understandability(features)


def create_prompt(text, prompt_es, prompt_ls, analysis_es, analysis_ls, analysis):
    """Create prompt and system message according the app settings."""
    if analysis:
        if leichte_sprache:
            final_prompt = analysis_ls.format(rules=RULES_LS, prompt=text)
            system = SYSTEM_MESSAGE_LS
        else:
            final_prompt = analysis_es.format(rules=RULES_ES, prompt=text)
            system = SYSTEM_MESSAGE_ES
    else:
        if leichte_sprache:
            if condense_text:
                final_prompt = prompt_ls.format(
                    rules=RULES_LS, completeness=REWRITE_CONDENSED, prompt=text
                )
            else:
                final_prompt = prompt_ls.format(
                    rules=RULES_LS, completeness=REWRITE_COMPLETE, prompt=text
                )
            system = SYSTEM_MESSAGE_LS
        else:
            final_prompt = prompt_es.format(
                rules=RULES_ES, completeness=REWRITE_COMPLETE, prompt=text
            )
            system = SYSTEM_MESSAGE_ES
    return final_prompt, system


def get_result_from_response(response):
    """Extract text between tags from response."""
    if leichte_sprache:
        result = re.findall(
            r"<leichtesprache>(.*?)</leichtesprache>", response, re.DOTALL
        )
    else:
        result = re.findall(
            r"<einfachesprache>(.*?)</einfachesprache>", response, re.DOTALL
        )
    result = "\n".join(result)
    return result.strip()


@st.cache_resource
# https://learn.microsoft.com/azure/cognitive-services/openai/how-to/create-resource?pivots=web-portal#create-a-resource
def get_azure_openai_client():
    return AzureOpenAI(
        api_key=AZURE_OPENAI_API_KEY,  # This is the default, but let's make it explicit.
        azure_endpoint=AZURE_OPENAI_ENDPOINT,
        # "2023-12-01-preview" or OPENAI_API_VERSION, see https://learn.microsoft.com/azure/ai-services/openai/reference#rest-api-versioning.
        api_version="2024-02-01",
    )


# Different shape than standard OpenAI API.
# See: https://github.com/openai/openai-python?tab=readme-ov-file#microsoft-azure-openai
# Changes compared to OpenAI library: https://learn.microsoft.com/en-us/azure/ai-services/openai/how-to/switching-endpoints


def invoke_azure_openai_model(
    text,
    deployment="simply-ch-gpt35turbo16k",
    temperature=TEMPERATURE,
    max_tokens=4096,
    analysis=False,
):
    """Invoke OpenAI model."""
    final_prompt, system = create_prompt(text, *OPENAI_TEMPLATES, analysis)
    try:
        message = azure_openai_client.chat.completions.create(
            model=deployment,  # Note that Azure OpenAI needs a deployment name (not the model name).
            temperature=temperature,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": final_prompt},
            ],
        )
        message = message.choices[0].message.content.strip()

        return True, get_result_from_response(message)
    except Exception as e:
        print(f"Error: {e}")
        return False, e


def enter_sample_text():
    """Enter sample text into the text input in the left column."""
    st.session_state.key_textinput = SAMPLE_TEXT_01


def create_download_link(text_input, response, analysis=False):
    """Create a downloadable Word document and download link of the results."""
    document = Document()

    h1 = document.add_heading("Ausgangstext")
    p1 = document.add_paragraph("\n" + text_input)

    if analysis:
        h2 = document.add_heading(f"Analyse von Sprachmodell {MODEL_NAME} von OpenAI")
    else:
        h2 = document.add_heading(
            f"Vereinfachter Text von Sprachmodell {MODEL_NAME} von OpenAI"
        )

    p2 = document.add_paragraph(response)

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer = document.sections[0].footer
    footer.paragraphs[
        0
    ].text = f"Erstellt am {timestamp} mit der Prototyp-App «Einfache Sprache», Statistisches Amt, Kanton Zürich.\nSprachmodell: {MODEL_NAME}\nVerarbeitungszeit: {time_processed:.1f} Sekunden"

    # Set font for all paragraphs.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_WORDDOC

    # Set font size for all headings.
    for paragraph in [h1, h2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_HEADING)

    # Set font size for all paragraphs.
    for paragraph in [p1, p2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_PARAGRAPH)

    # Set font and font size for footer.
    for run in footer.paragraphs[0].runs:
        run.font.name = "Arial"
        run.font.size = Pt(FONT_SIZE_FOOTER)

    section = document.sections[0]
    section.page_width = Inches(8.27)  # Width of A4 paper in inches
    section.page_height = Inches(11.69)  # Height of A4 paper in inches

    io_stream = io.BytesIO()
    document.save(io_stream)

    # # A download button unfortunately resets the app. So we use a link instead.
    # https://github.com/streamlit/streamlit/issues/4382#issuecomment-1223924851
    # https://discuss.streamlit.io/t/creating-a-pdf-file-generator/7613?u=volodymyr_holomb

    b64 = base64.b64encode(io_stream.getvalue())
    file_name = "Ergebnis.docx"

    caption = "Vereinfachten Text herunterladen"

    if analysis:
        file_name = "Analyse.docx"
        caption = "Analyse herunterladen"
    download_url = f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{file_name}">{caption}</a>'
    st.markdown(download_url, unsafe_allow_html=True)


def clean_log(text):
    """Remove linebreaks and tabs from log messages that otherwise would yield problems when parsing the logs."""
    text = text.replace("\n", " ")
    text = text.replace("\t", " ")
    return text


def log_event(
    text,
    response,
    do_analysis,
    do_simplification,
    leichte_sprache,
    time_processed,
    success,
):
    """Log event."""
    log_string = f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    log_string += f"\t{clean_log(text)}"
    log_string += f"\t{clean_log(response)}"
    log_string += f"\t{do_analysis}"
    log_string += f"\t{do_simplification}"
    log_string += f"\t{leichte_sprache}"
    log_string += f"\t{time_processed:.3f}"
    log_string += f"\t{success}"

    logging.warning(log_string)


# ---------------------------------------------------------------
# Main

azure_openai_client = get_azure_openai_client()

nlp = get_nlp_pipeline()
project_info = get_project_info()


st.markdown("## 🙋‍♀️ Sprache einfach vereinfachen (Version GPT-4o only)")
create_project_info(project_info)
st.caption(USER_WARNING, unsafe_allow_html=True)
st.markdown("---")

# Set up first row with all buttons and settings.
button_cols = st.columns([1, 1, 1, 2])
with button_cols[0]:
    st.button(
        "Beispiel einfügen",
        on_click=enter_sample_text,
        use_container_width=True,
        type="secondary",
        help="Fügt einen Beispieltext ein.",
    )
with button_cols[1]:
    do_analysis = st.button(
        "Text analysieren",
        use_container_width=True,
        help="Analysiert deinen Ausgangstext Satz für Satz.",
    )
with button_cols[2]:
    do_simplification = st.button(
        "Text vereinfachen",
        use_container_width=True,
        help="Vereinfacht deinen Ausgangstext.",
    )
with button_cols[3]:
    leichte_sprache = st.toggle(
        "Leichte Sprache",
        value=False,
        help="**Schalter aktiviert**: «Leichte Sprache». **Schalter nicht aktiviert**: «Einfache Sprache».",
    )
    if leichte_sprache:
        condense_text = st.toggle(
            "Text verdichten",
            value=True,
            help="**Schalter aktiviert**: Modell konzentriert sich auf essentielle Informationen und versucht, Unwichtiges wegzulassen. **Schalter nicht aktiviert**: Modell versucht, alle Informationen zu übernehmen.",
        )


# Instantiate empty containers for the text areas.
cols = st.columns([2, 2, 1])

with cols[0]:
    source_text = st.container()
with cols[1]:
    placeholder_result = st.empty()
with cols[2]:
    placeholder_analysis = st.empty()

# Populate containers.
with source_text:
    text_input = st.text_area(
        "Ausgangstext, den du vereinfachen möchtest",
        height=TEXT_AREA_HEIGHT,
        max_chars=MAX_CHARS_INPUT,
        key="key_textinput",
    )
with placeholder_result:
    text_output = st.text_area(
        "Ergebnis",
        height=TEXT_AREA_HEIGHT,
        key="key_textoutput",
    )
with placeholder_analysis:
    text_analysis = st.metric(
        label="Verständlichkeit 0-20 Punkte",
        value=None,
        delta=None,
        help="Texte in Einfacher Sprache haben meist einen Wert von 16 bis 20 Punkten, Texte in Leichter Sprache 18 bis 20 Punkte.",
    )


# Start processing if one of the processing buttons is clicked.
if do_simplification or do_analysis:
    start_time = time.time()
    if text_input == "":
        st.error("Bitte gib einen Text ein.")
        st.stop()

    score_source = get_understandability(text_input)
    score_source_rounded = int(np.round(score_source, 0))
    cefr_source = get_cefr_level(score_source)

    # Analyze source text and display results.
    with source_text:
        if score_source < LIMIT_HARD:
            st.markdown(
                f"Dein Ausgangstext ist **:red[schwer verständlich]**. ({score_source_rounded} von 20 Punkten). Das entspricht etwa dem **:red[Sprachniveau {cefr_source}]**."
            )
        elif score_source >= LIMIT_HARD and score_source < LIMIT_MEDIUM:
            st.markdown(
                f"Dein Ausgangstext ist **:orange[nur mässig verständlich]**. ({score_source_rounded} von 20 Punkten). Das entspricht etwa dem **:orange[Sprachniveau {cefr_source}]**."
            )
        else:
            st.markdown(
                f"Dein Ausgangstext ist **:green[gut verständlich]**. ({score_source_rounded} von 20 Punkten). Das entspricht etwa dem **:green[Sprachniveau {cefr_source}]**."
            )
        with placeholder_analysis.container():
            text_analysis = st.metric(
                label="Verständlichkeit 0-20 Punkte",
                value=score_source_rounded,
                delta=None,
                help="Verständlichkeit auf einer Skala von 0 bis 20 Punkten (von 0 = extrem schwer verständlich bis 20 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 16 bis 20 Punkten, Texte in Leichter Sprache 18 bis 20 Punkte.",
            )

        with placeholder_analysis.container():
            with st.spinner("Ich arbeite..."):
                # Regular text simplification or analysis.
                success, response = invoke_azure_openai_model(
                    text_input,
                    modelId=MODEL_CHOICE,
                    analysis=do_analysis,
                )

    if success is False:
        st.error(
            "Es ist ein Fehler bei der Abfrage der APIs aufgetreten. Bitte versuche es erneut. Alternativ überprüfe Code, API-Keys, Verfügbarkeit der Modelle und ggf. Internetverbindung."
        )
        time_processed = time.time() - start_time
        log_event(
            text_input,
            "Error from model call",
            do_analysis,
            do_simplification,
            leichte_sprache,
            time_processed,
            success,
        )

        st.stop()

    # Display results in UI.
    text = "Dein vereinfachter Text"
    if do_analysis:
        text = "Deine Analyse"
    # Often the models return the German letter «ß». Replace it with the Swiss «ss».
    response = response.replace("ß", "ss")
    time_processed = time.time() - start_time

    with placeholder_result.container():
        st.text_area(
            text,
            height=TEXT_AREA_HEIGHT,
            value=response,
        )
        if do_simplification:
            score_target = get_understandability(response)
            score_target_rounded = int(np.round(score_target, 0))
            cefr_target = get_cefr_level(score_target)
            if score_target < LIMIT_HARD:
                st.markdown(
                    f"Dein vereinfachter Text ist **:red[schwer verständlich]**. ({score_target_rounded} von 20 Punkten). Das entspricht etwa dem **:red[Sprachniveau {cefr_target}]**."
                )
            elif score_target >= LIMIT_HARD and score_target < LIMIT_MEDIUM:
                st.markdown(
                    f"Dein vereinfachter Text ist **:orange[nur mässig verständlich]**. ({score_target_rounded} von 20 Punkten). Das entspricht etwa dem **:orange[Sprachniveau {cefr_target}]**."
                )
            else:
                st.markdown(
                    f"Dein vereinfachter Text ist **:green[gut verständlich]**. ({score_target_rounded} von 20 Punkten). Das entspricht etwa dem **:green[Sprachniveau {cefr_target}]**."
                )
            with placeholder_analysis.container():
                text_analysis = st.metric(
                    label="Verständlichkeit 0-20 Punkte",
                    value=score_target_rounded,
                    delta=int(np.round(score_target - score_source, 0)),
                    help="Verständlichkeit auf einer Skala von 0 bis 20 Punkten (von 0 = extrem schwer verständlich bis 20 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 16 bis 20 Punkten.",
                )

                create_download_link(text_input, response)
                st.caption(f"Verarbeitet in {time_processed:.1f} Sekunden.")
        else:
            with placeholder_analysis.container():
                text_analysis = st.metric(
                    label="Verständlichkeit 0-20 Punkte",
                    value=score_source_rounded,
                    help="Verständlichkeit auf einer Skala von 0 bis 20 Punkten (von 0 = extrem schwer verständlich bis 20 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 16 bis 20 Punkten.",
                )
                create_download_link(text_input, response, analysis=True)
                st.caption(f"Verarbeitet in {time_processed:.1f} Sekunden.")

        log_event(
            text_input,
            response,
            do_analysis,
            do_simplification,
            leichte_sprache,
            time_processed,
            success,
        )
        st.stop()
